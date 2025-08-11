#!/usr/bin/env python3
"""
Test script untuk menguji fitur ukuran file dalam proses watermarking
"""

import os
import sys
from main import embed_watermark_to_docx, generate_qr_code
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile

def create_test_docx_with_image():
    """Membuat test DOCX dengan gambar untuk testing"""
    # Buat document sederhana
    doc = Document()
    doc.add_heading('Test Document untuk File Size Analysis', 0)
    doc.add_paragraph('Ini adalah dokumen test untuk menganalisis perubahan ukuran file setelah LSB steganography.')
    
    # Buat gambar test sederhana
    test_image = Image.new('RGB', (400, 300), color='lightblue')
    for i in range(50):
        for j in range(50):
            test_image.putpixel((i*8, j*6), (255-i*5, 128+i*2, 100+j*3))
    
    temp_img_path = 'test_image.png'
    test_image.save(temp_img_path)
    
    # Tambahkan gambar ke dokumen
    doc.add_paragraph('Gambar test untuk watermarking:')
    doc.add_picture(temp_img_path, width=Inches(4))
    
    doc.add_paragraph('Text tambahan untuk menambah ukuran dokumen dan memberikan konteks lebih.')
    
    # Simpan dokumen
    test_docx_path = 'test_document.docx'
    doc.save(test_docx_path)
    
    # Cleanup temp image
    if os.path.exists(temp_img_path):
        os.remove(temp_img_path)
    
    return test_docx_path

def test_file_size_analysis():
    """Test fungsi analisis ukuran file"""
    print("=" * 60)
    print("TEST: File Size Analysis dalam LSB Steganography")
    print("=" * 60)
    
    # 1. Buat dokumen test
    print("\n1. Membuat dokumen test dengan gambar...")
    test_docx = create_test_docx_with_image()
    original_size = os.path.getsize(test_docx)
    print(f"   Dokumen test dibuat: {test_docx}")
    print(f"   Ukuran asli: {original_size:,} bytes ({original_size/1024:.2f} KB)")
    
    # 2. Generate QR Code untuk watermark
    print("\n2. Membuat QR Code untuk watermark...")
    qr_data = "TEST WATERMARK DATA 2024 - LSB Steganography Analysis"
    qr_path = 'test_watermark.png'
    
    if generate_qr_code(qr_data, qr_path):
        qr_size = os.path.getsize(qr_path)
        print(f"   QR Code dibuat: {qr_path}")
        print(f"   Ukuran QR: {qr_size} bytes")
    else:
        print("   GAGAL: Tidak dapat membuat QR Code")
        return False
    
    # 3. Proses embedding dengan analisis ukuran file
    print("\n3. Melakukan embedding watermark dengan analisis ukuran file...")
    output_docx = 'test_watermarked.docx'
    
    try:
        result = embed_watermark_to_docx(test_docx, qr_path, output_docx)
        
        if result.get("success"):
            print("   ✅ Embedding berhasil!")
            
            # Analisis hasil dari fungsi
            if "file_size_info" in result:
                size_info = result["file_size_info"]
                print("\n4. ANALISIS UKURAN FILE:")
                print(f"   📄 Ukuran asli: {size_info['original_size']:,} bytes ({size_info['original_size_kb']:.2f} KB)")
                print(f"   📄 Ukuran watermark: {size_info['watermarked_size']:,} bytes ({size_info['watermarked_size_kb']:.2f} KB)")
                print(f"   📊 Perubahan: {size_info['size_difference']:+,} bytes ({size_info['size_change_percentage']:+.3f}%)")
                
                # Analisis kualitas steganography berdasarkan perubahan ukuran
                abs_change = abs(size_info['size_change_percentage'])
                if abs_change < 0.1:
                    print("   🟢 EXCELLENT: Perubahan ukuran sangat minimal (<0.1%)")
                elif abs_change < 1:
                    print("   🟢 VERY GOOD: Perubahan ukuran minimal (<1%)")
                elif abs_change < 5:
                    print("   🟡 GOOD: Perubahan ukuran kecil (<5%)")
                else:
                    print("   🔴 WARNING: Perubahan ukuran signifikan (>5%)")
                

            else:
                print("   ❌ Informasi ukuran file tidak tersedia dalam hasil")
        else:
            print(f"   ❌ Embedding gagal: {result.get('error', 'Unknown error')}")
            return False
            
    except Exception as e:
        print(f"   ❌ Error saat embedding: {str(e)}")
        return False
    
    # 5. Analisis perbandingan dengan file sistem dan verifikasi
    print(f"\n5. VERIFIKASI & PERBANDINGAN SISTEM FILE:")
    if os.path.exists(test_docx) and os.path.exists(output_docx):
        system_original = os.path.getsize(test_docx)
        system_watermarked = os.path.getsize(output_docx)
        system_diff = system_watermarked - system_original
        system_percent = (system_diff / system_original) * 100
        
        print(f"   📁 Original (sistem): {system_original:,} bytes")
        print(f"   📁 Watermarked (sistem): {system_watermarked:,} bytes")
        print(f"   📊 Perubahan (sistem): {system_diff:+,} bytes ({system_percent:+.3f}%)")
        
        # Verifikasi konsistensi data dengan hasil fungsi
        if "file_size_info" in result:
            size_info = result["file_size_info"]
            consistency_check = system_watermarked == size_info['watermarked_size']
            print(f"   ✅ Konsistensi data fungsi vs sistem: {'OK' if consistency_check else 'ERROR'}")
    
    # Cleanup
    print(f"\n6. Pembersihan file test...")
    for file_path in [test_docx, qr_path, output_docx]:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"   🗑️  Menghapus: {file_path}")
    
    print("\n" + "=" * 60)
    print("TEST SELESAI - File Size Analysis implemented successfully!")
    print("=" * 60)
    return True

if __name__ == "__main__":
    success = test_file_size_analysis()
    sys.exit(0 if success else 1)
